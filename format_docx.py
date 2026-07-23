import os
import docx
from docx.shared import Pt, Inches
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn

def add_background_to_first_page(doc, image_path):
    if not os.path.exists(image_path):
        print(f"Cover background image not found at: {image_path}")
        return
        
    section = doc.sections[0]
    section.different_first_page_header_footer = True
    
    first_page_header = section.first_page_header
    for p in list(first_page_header.paragraphs):
        p_elm = p._element
        p_elm.getparent().remove(p_elm)
        
    p = first_page_header.add_paragraph()
    p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    
    run = p.add_run()
    inline_shape = run.add_picture(image_path, width=Inches(8.27), height=Inches(11.69))
    
    inline = inline_shape._inline
    graphic = inline.graphic
    
    cx = int(8.27 * 914400)
    cy = int(11.69 * 914400)
    
    anchor_xml = f'''
    <wp:anchor xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
               distT="0" distB="0" distL="0" distR="0" simplePos="0" relativeHeight="251658240"
               behindDoc="1" locked="0" layoutInCell="1" allowOverlap="1">
        <wp:simplePos x="0" y="0"/>
        <wp:positionH relativeFrom="page">
            <wp:posOffset>0</wp:posOffset>
        </wp:positionH>
        <wp:positionV relativeFrom="page">
            <wp:posOffset>0</wp:posOffset>
        </wp:positionV>
        <wp:extent cx="{cx}" cy="{cy}"/>
        <wp:effectExtent l="0" t="0" r="0" b="0"/>
        <wp:wrapNone/>
        <wp:docPr id="999" name="Cover Background"/>
        <wp:cNvGraphicFramePr/>
    </wp:anchor>
    '''
    anchor = parse_xml(anchor_xml)
    anchor.append(graphic)
    inline.getparent().replace(inline, anchor)
    print("  Successfully added cover background image")

def format_document(file_path):
    print(f"Loading document: {file_path}")
    doc = docx.Document(file_path)

    # 1. Set A4 Size (21cm x 29.7cm) & Margins (1 inch / 2.54 cm) for all sections & remove simple borders
    for section in doc.sections:
        section.page_width = Inches(8.27)   # A4 width
        section.page_height = Inches(11.69) # A4 height
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

        # Remove page borders (bingkai)
        sec_pr = section._sectPr
        for pg_border in list(sec_pr.xpath('w:pgBorders')):
            sec_pr.remove(pg_border)

    # Apply coffee-themed cover background to the first page (cover page)
    add_background_to_first_page(doc, "c:/Users/Thinkpad/kopikeliling/public/coffee_cover_bg.png")

    # 2. Modify default style fonts and line spacing
    # Normal style (Body Text)
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Times New Roman'
    style_normal.font.size = Pt(12)
    style_normal.paragraph_format.line_spacing = 1.15

    # Modify heading styles to also use Times New Roman
    for style_name in ['Heading 1', 'Heading 2', 'Heading 3', 'Title', 'Subtitle', 'List Paragraph']:
        if style_name in doc.styles:
            s = doc.styles[style_name]
            s.font.name = 'Times New Roman'
            if style_name == 'Heading 1':
                s.font.size = Pt(16)
                s.font.bold = True
            elif style_name == 'Heading 2':
                s.font.size = Pt(14)
                s.font.bold = True
            elif style_name == 'Heading 3':
                s.font.size = Pt(12)
                s.font.bold = True
            elif style_name == 'Title':
                s.font.size = Pt(24)
                s.font.bold = True
            elif style_name == 'List Paragraph':
                s.font.size = Pt(12)
                s.paragraph_format.line_spacing = 1.15

    # 3. Apply Times New Roman and spacing 1.15 to all paragraphs and runs explicitly
    for paragraph in doc.paragraphs:
        paragraph.paragraph_format.line_spacing = 1.15
        
        # Don't override title or heading sizes, but force font family to Times New Roman
        for run in paragraph.runs:
            run.font.name = 'Times New Roman'
            
            # Apply XML element settings to ensure word doesn't fallback to Calibri
            rPr = run._r.get_or_add_rPr()
            rFonts = OxmlElement('w:rFonts')
            rFonts.set(qn('w:ascii'), 'Times New Roman')
            rFonts.set(qn('w:hAnsi'), 'Times New Roman')
            rFonts.set(qn('w:cs'), 'Times New Roman')
            rFonts.set(qn('w:eastAsia'), 'Times New Roman')
            rPr.append(rFonts)

            # Ensure run has font size set
            if run.font.size is None:
                if not paragraph.style.name.startswith('Heading') and paragraph.style.name != 'Title':
                    run.font.size = Pt(12)

    # 4. Spacing and font in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.line_spacing = 1.15
                    for run in paragraph.runs:
                        run.font.name = 'Times New Roman'
                        
                        rPr = run._r.get_or_add_rPr()
                        rFonts = OxmlElement('w:rFonts')
                        rFonts.set(qn('w:ascii'), 'Times New Roman')
                        rFonts.set(qn('w:hAnsi'), 'Times New Roman')
                        rPr.append(rFonts)
                        
                        if run.font.size is None:
                            run.font.size = Pt(11) # Table text can be 11pt

    # Save changes
    doc.save(file_path)
    print(f"Successfully formatted: {file_path}")

if __name__ == "__main__":
    format_document('c:/Users/Thinkpad/kopikeliling/BisnisTIK_KMIPN2026_KopiKuy.docx')
